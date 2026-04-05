import os
import base64
import json
import re
import io

from fastapi import FastAPI, Header, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import google.generativeai as genai

load_dotenv()

app = FastAPI(title="AI Document Analyzer", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
MY_API_KEY = os.getenv("MY_API_KEY", "sk_track2_987654321")

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")


class DocRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str


# ─── Text Extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {e}")


def extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += "\n" + row_text
        return text
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {e}")


def extract_text_from_image(data: bytes) -> str:
    try:
        import base64 as b64
        image_b64 = b64.b64encode(data).decode('utf-8')

        if data[:4] == b'\x89PNG':
            mime = 'image/png'
        elif data[:2] == b'\xff\xd8':
            mime = 'image/jpeg'
        else:
            mime = 'image/jpeg'

        response = model.generate_content([
            {
                "inline_data": {
                    "mime_type": mime,
                    "data": image_b64
                }
            },
            "Extract ALL text from this image exactly as it appears. Return only the raw text, nothing else."
        ])
        return response.text.strip()
    except Exception as e:
        raise ValueError(f"Image extraction failed: {e}")


def extract_text(file_type: str, b64: str) -> str:
    data = base64.b64decode(b64)
    ft = file_type.lower().strip()
    if ft == "pdf":
        return extract_text_from_pdf(data)
    elif ft == "docx":
        return extract_text_from_docx(data)
    elif ft in ("image", "png", "jpg", "jpeg", "bmp", "tiff", "gif"):
        return extract_text_from_image(data)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ─── AI Analysis ──────────────────────────────────────────────────────────────

def analyze_with_gemini(text: str) -> dict:
    prompt = f"""You are an expert document analyzer. Read the document text below carefully and extract information.

Return ONLY a valid JSON object. No markdown, no explanation, no code fences.

JSON structure:
{{
  "summary": "Write 2-3 sentences summarizing what this document is about, who it involves, and its key purpose.",
  "entities": {{
    "names": ["Extract ONLY actual human person names (first name + last name). Do NOT include job titles, company names, or generic terms."],
    "dates": ["Extract ALL dates mentioned in any format like DD/MM/YYYY, Month DD YYYY, DD Month YYYY, etc."],
    "organizations": ["Extract ONLY real company names, organization names, institution names. Do NOT include generic words."],
    "amounts": ["Extract ALL monetary amounts with their currency symbols like Rs., ₹, $, USD, etc."]
  }},
  "sentiment": "Classify the overall tone as exactly one of: Positive, Neutral, Negative"
}}

IMPORTANT RULES:
- names: ONLY real human names (e.g. 'Ravi Kumar', 'John Smith'). NOT 'Artificial Intelligence', NOT job titles.
- dates: Include ALL date formats found.
- organizations: ONLY real org names (e.g. 'ABC Pvt Ltd', 'Google', 'HDFC Bank'). NOT generic words.
- amounts: Include currency symbol (e.g. '₹10,000', '$500', 'Rs. 2500').
- If nothing found for a field, return empty list [].
- sentiment: invoices/reports = Neutral, positive news = Positive, complaints/negative = Negative.

Document text:
{text[:5000]}"""

    response = model.generate_content(
        prompt,
        generation_config={"temperature": 0.1}
    )

    raw = response.text.strip()
    raw = re.sub(r"```json|```", "", raw).strip()

    match = re.search(r'\{.*\}', raw, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass

    raise ValueError("Gemini did not return valid JSON")


# ─── API Routes ───────────────────────────────────────────────────────────────

@app.get("/")
def root():
    return {
        "message": "AI Document Analyzer API is running!",
        "docs": "/docs",
        "powered_by": "Google Gemini 2.5 Flash"
    }


@app.get("/api/document-analyze")
def analyze_get():
    return {
        "message": "Use POST method to analyze a document.",
        "endpoint": "POST /api/document-analyze",
        "required_headers": ["Content-Type: application/json", "x-api-key: YOUR_KEY"],
        "body": {"fileName": "doc.pdf", "fileType": "pdf", "fileBase64": "..."}
    }


@app.post("/api/document-analyze")
async def analyze_document(request: Request, x_api_key: str = Header(None)):
    # Strict API key check
    if x_api_key != MY_API_KEY:
        raise HTTPException(
            status_code=401,
            detail="Unauthorized — invalid or missing x-api-key header"
        )

    try:
        data = await request.json()
        fileName = data.get('fileName', '')
        fileType = data.get('fileType', '')
        fileBase64 = data.get('fileBase64', '')

        # Step 1: Extract text
        text = extract_text(fileType, fileBase64)

        if not text or not text.strip():
            return {
                "status": "error",
                "fileName": fileName,
                "message": "No text could be extracted from the document."
            }

        # Step 2: Analyze with Gemini
        result = analyze_with_gemini(text)

        return {
            "status": "success",
            "fileName": fileName,
            "summary": result.get("summary", ""),
            "entities": {
                "names": result.get("entities", {}).get("names", []),
                "dates": result.get("entities", {}).get("dates", []),
                "organizations": result.get("entities", {}).get("organizations", []),
                "amounts": result.get("entities", {}).get("amounts", [])
            },
            "sentiment": result.get("sentiment", "Neutral")
        }

    except HTTPException:
        raise
    except Exception as e:
        return {
            "status": "error",
            "fileName": data.get('fileName', ''),
            "message": str(e)
        }


handler = app